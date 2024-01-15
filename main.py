import streamlit as st
from streamlit_option_menu import option_menu
from io import BytesIO
from PIL import Image
import docx2txt
from docx import Document
import base64
import string
import mysql.connector

#Database Connection
db_connection = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="steganografi"
)

db_cursor = db_connection.cursor()

png_file_path =""
word_file_path = ""

# machine one
def embed_text_to_image(text, output_path):
    global png_file_path, word_file_path

    # png area
    def rgb_to_binary(rgb):
        return ''.join(format(value, '08b') for value in rgb)

    def extracting_image_rgb(img_path):
        image = Image.open(img_path)
        pixel_values = list(image.getdata())
        binary_pixel_values = [rgb_to_binary(rgb) for rgb in pixel_values]
        image.close()
        bit = []
        k = ""
        i = 0

        for three_binary in binary_pixel_values:
            for binary in three_binary:
                i+=1
                k+=binary
                if i == 8 or i == 16 or i == 24:
                    bit.append(k)
                    k = ""
                    i = 0
        return bit

    def get_png_resolution(image_path):
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                return width, height
        except Exception as e:
            print(f"Error: {e}")
            return None

    png_binary = extracting_image_rgb(png_file_path)
    resolution = get_png_resolution(png_file_path)

    # word area
    def read_docx(file_path):
        doc = Document(file_path)

        # Initialize an empty string to store the text
        text_content = ""

        # Iterate through paragraphs and add text to the string
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"

        return text_content

    def text_to_binary(text):
        binary_data = ''.join(format(ord(char), '08b') for char in text)
        return binary_data

    word_text = read_docx(text)
    word_text = ''.join(char for char in word_text if char in string.printable)
    word_binary = text_to_binary(word_text)

    # checking the limit
    length_count_png_binary = len(png_binary) * 3
    length_count_word_binary = len(word_binary) * 8
    length_count_word_binary = length_count_word_binary + 32

    dum_list = []
    for i in word_binary:
        dum_list.append(i)
    word_binary = dum_list

    dum_list = []
    decimal_number = len(word_binary)
    binary_string = bin(decimal_number)[2:]

    dum_list = []
    j = 0
    for i in range(0, 32):
        if i < (32-len(binary_string)):
            dum_list.append('0')
            continue
        dum_list.append(binary_string[j])
        j = j + 1
    count_word = dum_list

    word_binary = count_word + word_binary

    # steganografi program
    new_png_binary = []
    length_png = len(png_binary)
    length_word = len(word_binary)
    for i in range(length_word):
        str_png_binary = png_binary[i]
        new_binary = str_png_binary[:-1] + word_binary[i]
        new_png_binary.append(new_binary)

    for i in range(length_word, length_png):
        new_png_binary.append(png_binary[i])

        #  Code for construct binary png to png
    def create_image_from_pixel_data(pixel_data, width, height):
        # Create a new image with the specified width and height
        image = Image.new('RGB', (width, height))

        # Create a pixel access object
        pixels = image.load()

        # Set pixel values from the pixel data
        for y in range(height):
            for x in range(width):
                pixel_index = y * width * 3 + x * 3
                pixels[x, y] = (pixel_data[pixel_index], pixel_data[pixel_index + 1], pixel_data[pixel_index + 2])

        return image

    rgb_values = []
    dum_list = []

    for i in range(0, len(new_png_binary), 3):
        dum_list = []
        for j in range(3):
            dum_list.append(int(new_png_binary[i+j],2))
        rgb_values.append(dum_list)

    pixel_data = [value for rgb_list in rgb_values for value in rgb_list]

    image = create_image_from_pixel_data(pixel_data, resolution[0], resolution[1])
    image.save(output_path)

def checking_limit(text):
    global png_file_path, word_file_path

    # png area
    def rgb_to_binary(rgb):
        return ''.join(format(value, '08b') for value in rgb)

    def extracting_image_rgb(img_path):
        image = Image.open(img_path)
        pixel_values = list(image.getdata())
        binary_pixel_values = [rgb_to_binary(rgb) for rgb in pixel_values]
        image.close()
        bit = []
        k = ""
        i = 0

        for three_binary in binary_pixel_values:
            for binary in three_binary:
                i+=1
                k+=binary
                if i == 8 or i == 16 or i == 24:
                    bit.append(k)
                    k = ""
                    i = 0
        return bit

    def get_png_resolution(image_path):
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                return width, height
        except Exception as e:
            print(f"Error: {e}")
            return None

    png_binary = extracting_image_rgb(png_file_path)
    resolution = get_png_resolution(png_file_path)

    # word area
    def read_docx(file_path):
        doc = Document(file_path)

        # Initialize an empty string to store the text
        text_content = ""

        # Iterate through paragraphs and add text to the string
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"

        return text_content

    def text_to_binary(text):
        binary_data = ''.join(format(ord(char), '08b') for char in text)
        return binary_data

    word_text = read_docx(text)
    word_text = ''.join(char for char in word_text if char in string.printable)
    word_binary = text_to_binary(word_text)

    # checking the limit
    length_count_png_binary = len(png_binary) * 3
    length_count_word_binary = len(word_binary) * 8
    length_count_word_binary = length_count_word_binary + 32

    output = 0
    if length_count_png_binary < length_count_word_binary:
        output = 1

    return output

# machine two
# 2nd pages
def embed_image_to_text(image_path):
    def rgb_to_binary(rgb):
        return ''.join(format(value, '08b') for value in rgb)

    def extracting_image_rgb(image):
        try:
            if isinstance(image, Image.Image):
                pixel_values = list(image.getdata())
                binary_pixel_values = [rgb_to_binary(rgb) for rgb in pixel_values]

                bit = []
                k = ""
                i = 0

                for three_binary in binary_pixel_values:
                    for binary in three_binary:
                        i+=1
                        k+=binary
                        if i == 8 or i == 16 or i == 24:
                            bit.append(k)
                            k = ""
                            i = 0
                return bit
            else:
                raise ValueError("Objek gambar tidak valid")
        except Exception as e:
            print(f"Error: {e}")
            return None

    png_binary = extracting_image_rgb(image_path)

    count_binary = ""
    for i in range(0, 32):
        str_png_binary = png_binary[i]
        count_binary+=(str_png_binary[-1])

    binary_string = count_binary
    count_decimal = int(binary_string, 2)

    # Code to get every LSB and put it together
    word_binary = ""
    for i in range(32, count_decimal+32):
        str_png_binary = png_binary[i]
        word_binary+=(str_png_binary[-1])

    def binary_to_text(binary_string):
        text = ''.join(chr(int(binary_string[i:i+8], 2)) for i in range(0, len(binary_string), 8))
        return text

    # Example usage
    return binary_to_text(word_binary)

# main
st.set_page_config(
        page_title="STEGO~",
        page_icon=":key",
        initial_sidebar_state="auto",
)

with st.sidebar:
    selected = option_menu("Main Menu", ["Home", "Halaman 1", "Halaman 2"],
    icons=['house', 'gear', 'key'], menu_icon="cast", default_index=0)
if(selected=="Home"):
    st.title("About Steganografi Paper")
    st.caption("Ini adalah halaman utama web")
    st.text_area("Definisi Steganografi",
                    "Menurut Munir (2009), Steganografi adalah ilmu dan seni menyembunyikan"
                    "pesan rahasia (hiding message) sedemikian sehingga keberadaan (eksistensi)"
                    "pesan yang tidak terdeteksi oleh indera manusia." )
    st.subheader("Cara Kerja")
    st.markdown("***Cara Kerja Steganografi***")
    st.markdown("1. Input Image dan Input Text")
    st.markdown("2. Program dieksekusi")
    st.markdown("3. Output Image yang mempunyai pesan tersembunyi didalamnya")
    st.markdown("***Cara Program Mengeksekusi***")
    st.markdown("1. Mengekstrak setiap bit dari Image dan Mengekstrak bit dari test")
    st.markdown("2. Menggunakan LSB untuk menyisipkan setiap digit bit text ke Image")
    st.markdown("3. Menciptakan kembali Image yang telah disisipi bit di akhir bitnya")

elif(selected=="Halaman 1"):
    st.title("Halaman 1")
    st.markdown("Penyisipan setiap isi teks pada file DOCX (dalam bentuk binary) ke dalam setiap pixel pada file PNG atau JPG")
    st.markdown("***Input File PNG atau JPG dan File DOCX***")

    # Membuat komponen untuk mengunggah file dengan dukungan untuk beberapa file.
    uploaded_image = st.file_uploader("Pilih file Foto", type=["jpg", "png"], accept_multiple_files=False)
    uploaded_text = st.file_uploader("Pilih file DOCX", type=["docx"], accept_multiple_files=False)

    if st.button("Submit"):
        if uploaded_image is not None and uploaded_text is not None:
            try:
                # Save image and text to temporary files
                temp_image_path = "temp_image.jpg" or "temp_image.png"
                temp_text_path = "temp_text.docx"
                png_file_path = "temp_image.jpg"
                word_file_path ="temp_text.docx"
                with open(temp_image_path, "wb") as temp_image_file:
                    temp_image_file.write(uploaded_image.read())

                with open(temp_text_path, "wb") as temp_text_file:
                    temp_text_file.write(uploaded_text.read())

                # Get image resolution
                with Image.open(temp_image_path) as img:
                    width, height = img.size

                boolean_1stmachine = checking_limit(temp_text_path)
                if boolean_1stmachine == 1:
                    st.write("please use bigger image")
                else:
                    # Embed text to image
                    output_image_path = "output_image.png"
                    embed_text_to_image(temp_text_path, output_image_path)

                    db_cursor.execute('''
                        INSERT INTO embedded_images (image_name, text_name, output_image_path)
                                    VALUES (%s, %s, %s)
                                    ''', (uploaded_image.name, uploaded_text.name, output_image_path))

                    db_connection.commit()

                    # Display the output image
                    st.image(output_image_path, caption="Gambar yang telah disisipi teks", use_column_width=True)
                    # Download button for the output image
                    output_image_download_label = "Download Output Image"
                    output_image_base64 = base64.b64encode(open(output_image_path, "rb").read()).decode()
                    href = f'<a href="data:file/png;base64,{output_image_base64}" download="output_image.png">{output_image_download_label}</a>'
                    st.markdown(href, unsafe_allow_html=True)                
            except Exception as e:
                st.error(f"Terjadi kesalahan: {e}")
            finally:
                db_cursor.close()
                db_connection.close()
        else:
            st.write("Silakan pilih file gambar (JPG or PNG) dan file teks (DOCX) sebelum menekan tombol Submit.")  
    def process_uploaded_file(uploaded_file):
        # Memeriksa apakah file telah diunggah
        if uploaded_file is not None:
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  # Mengecek jika file adalah docx
                # Menggunakan docx2txt untuk mengonversi dokumen Word ke teks
                text_data = docx2txt.process(uploaded_file)
                st.write("Data teks dari dokumen Word (docx):")
                st.write(text_data)
            elif uploaded_file.type == "image/jpeg" or uploaded_file.type == "image/jpg" or uploaded_file.type == "image/png": # Mengecek jika file adalah jpg
                # Menggunakan PIL untuk memuat dan menampilkan gambar JPG
                image = Image.open(uploaded_file)
                st.image(image, caption="Gambar yang diunggah", use_column_width=True)
            else:
                st.write("Tipe file tidak didukung. Harap pilih file dengan tipe docx atau jpg.")

elif(selected=="Halaman 2"):
    st.title("Halaman 2")
    st.markdown("Mengekstrak kembali isi dari file PNG atau JPG yang telah disisipkan")
    st.markdown("***Input File PNG atau JPG yang telah diunduh dari halaman 1***")
    uploaded_file = st.file_uploader("Pilih file JPG or PNG", type=["jpg","png"], accept_multiple_files=False)
    try:
        if st.button("Submit"):
            if uploaded_file is not None:
                if uploaded_file.type != "image/png":
                    st.write("Tipe file tidak didukung. Harap pilih file dengan tipe .png.")
                else:
                    image = Image.open(uploaded_file)
                    st.image(image, caption="Gambar yang diunggah", use_column_width=True)

                    text = embed_image_to_text(image)

                    #Insert data into Mysql Database
                    db_cursor.execute('''
                        INSERT INTO embedded_files (output_image_name, output_file_path)
                        VALUES (%s, %s)
                    ''', (uploaded_file.name, text))

                    db_connection.commit()

                    def write_text_to_word(text, outputh_path='output.docx'):
                        doc = Document()
                        doc.add_paragraph(text)

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        return buffer

                    buffer = write_text_to_word(text)
                    st.download_button(
                        label="Download Word Document",
                        data=buffer,
                        file_name='output.docx',
                        key="download_button"

                    )
    except Exception as e:
        st.error(f"Terjadi kesalahan: {e}")
    finally:
        db_cursor.close()
        db_connection.close()